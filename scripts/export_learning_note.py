from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
)


def iter_blocks(markdown: str):
    in_code = False
    code_lines: list[str] = []
    paragraph_lines: list[str] = []

    def flush_paragraph():
        nonlocal paragraph_lines
        if paragraph_lines:
            yield ("paragraph", " ".join(line.strip() for line in paragraph_lines))
            paragraph_lines = []

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if line.startswith("```"):
            if in_code:
                yield ("code", "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                yield from flush_paragraph()
                in_code = True
            continue

        if in_code:
            code_lines.append(raw_line)
            continue

        if not line.strip() or line.strip() == ">":
            yield from flush_paragraph()
            continue

        if line.startswith("---"):
            yield from flush_paragraph()
            yield ("rule", "")
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading_match:
            yield from flush_paragraph()
            yield (f"h{len(heading_match.group(1))}", heading_match.group(2).strip())
            continue

        if line.startswith("> "):
            yield from flush_paragraph()
            yield ("quote", line[2:].strip())
            continue

        if re.match(r"^\s*[-*]\s+", line):
            yield from flush_paragraph()
            yield ("bullet", re.sub(r"^\s*[-*]\s+", "", line).strip())
            continue

        if re.match(r"^\s*\d+\.\s+", line):
            yield from flush_paragraph()
            yield ("number", re.sub(r"^\s*\d+\.\s+", "", line).strip())
            continue

        paragraph_lines.append(line)

    if in_code and code_lines:
        yield ("code", "\n".join(code_lines))
    yield from flush_paragraph()


def strip_inline_markup(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text


DOCX_BODY_FONT = "Arial Unicode MS"
DOCX_CODE_FONT = "Menlo"


def set_run_font(run, font_name: str, size: Pt | None = None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = size


def set_style_font(style, font_name: str, size: Pt | None = None):
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        style.font.size = size


def export_docx(blocks, output: Path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    set_style_font(styles["Normal"], DOCX_BODY_FONT, Pt(10.5))

    for kind, text in blocks:
        text = strip_inline_markup(text)
        if kind == "h1":
            p = doc.add_heading(text, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, DOCX_BODY_FONT)
        elif kind.startswith("h"):
            level = min(int(kind[1:]), 4)
            p = doc.add_heading(text, level=level)
            for run in p.runs:
                set_run_font(run, DOCX_BODY_FONT)
        elif kind == "bullet":
            p = doc.add_paragraph(text, style="List Bullet")
            for run in p.runs:
                set_run_font(run, DOCX_BODY_FONT)
        elif kind == "number":
            p = doc.add_paragraph(text, style="List Number")
            for run in p.runs:
                set_run_font(run, DOCX_BODY_FONT)
        elif kind == "quote":
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.italic = True
                set_run_font(run, DOCX_BODY_FONT)
        elif kind == "code":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(text)
            set_run_font(run, DOCX_CODE_FONT, Pt(8.5))
        elif kind == "rule":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
        else:
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                set_run_font(run, DOCX_BODY_FONT)

    doc.save(output)


def register_pdf_font() -> str:
    candidates = [
        Path("/Library/Fonts/Arial Unicode.ttf"),
        Path("/System/Library/Fonts/STHeiti Medium.ttc"),
        Path("/System/Library/Fonts/Hiragino Sans GB.ttc"),
    ]
    for font in candidates:
        if font.exists():
            pdfmetrics.registerFont(TTFont("DocFont", str(font)))
            return "DocFont"
    return "Helvetica"


def xml_escape(text: str) -> str:
    return (
        strip_inline_markup(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def export_pdf(blocks, output: Path, title: str):
    font_name = register_pdf_font()
    styles = getSampleStyleSheet()
    base = ParagraphStyle(
        "BaseCN",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=10,
        leading=15,
        alignment=TA_LEFT,
        spaceAfter=6,
    )
    h1 = ParagraphStyle(
        "H1CN",
        parent=base,
        fontSize=20,
        leading=26,
        textColor=colors.HexColor("#1f2937"),
        spaceAfter=18,
    )
    h2 = ParagraphStyle(
        "H2CN",
        parent=base,
        fontSize=15,
        leading=20,
        textColor=colors.HexColor("#111827"),
        spaceBefore=12,
        spaceAfter=8,
    )
    h3 = ParagraphStyle(
        "H3CN",
        parent=base,
        fontSize=12.5,
        leading=17,
        textColor=colors.HexColor("#374151"),
        spaceBefore=8,
        spaceAfter=6,
    )
    quote = ParagraphStyle(
        "QuoteCN",
        parent=base,
        leftIndent=14,
        textColor=colors.HexColor("#4b5563"),
        italic=True,
    )
    bullet = ParagraphStyle(
        "BulletCN",
        parent=base,
        leftIndent=14,
        firstLineIndent=-8,
    )
    code = ParagraphStyle(
        "CodeCN",
        parent=base,
        fontName=font_name,
        fontSize=8.2,
        leading=11,
        leftIndent=10,
        backColor=colors.HexColor("#f3f4f6"),
        borderPadding=6,
        spaceBefore=4,
        spaceAfter=8,
    )

    story = []
    for kind, text in blocks:
        if kind == "h1":
            story.append(Paragraph(xml_escape(text), h1))
        elif kind == "h2":
            story.append(Paragraph(xml_escape(text), h2))
        elif kind.startswith("h"):
            story.append(Paragraph(xml_escape(text), h3))
        elif kind == "bullet":
            story.append(Paragraph(f"- {xml_escape(text)}", bullet))
        elif kind == "number":
            story.append(Paragraph(xml_escape(text), bullet))
        elif kind == "quote":
            story.append(Paragraph(xml_escape(text), quote))
        elif kind == "code":
            story.append(Preformatted(text, code))
        elif kind == "rule":
            story.append(Spacer(1, 0.2 * cm))
        else:
            story.append(Paragraph(xml_escape(text), base))

    doc = SimpleDocTemplate(
        str(output),
        pagesize=A4,
        rightMargin=1.7 * cm,
        leftMargin=1.7 * cm,
        topMargin=1.6 * cm,
        bottomMargin=1.6 * cm,
        title=title,
    )
    doc.build(story)


def main():
    if len(sys.argv) != 4:
        raise SystemExit(
            "Usage: export_learning_note.py input.md output.docx output.pdf"
        )

    input_path = Path(sys.argv[1])
    docx_path = Path(sys.argv[2])
    pdf_path = Path(sys.argv[3])
    markdown = input_path.read_text(encoding="utf-8")
    blocks = list(iter_blocks(markdown))
    title = next((text for kind, text in blocks if kind == "h1"), input_path.stem)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    pdf_path.parent.mkdir(parents=True, exist_ok=True)

    export_docx(blocks, docx_path)
    export_pdf(blocks, pdf_path, title)


if __name__ == "__main__":
    main()
